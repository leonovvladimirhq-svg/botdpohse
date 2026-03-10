import os
import csv
import logging
from pathlib import Path
from datetime import datetime

from lxml import etree
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    filters, ContextTypes, ConversationHandler,
)

# --- Настройки ---
load_dotenv()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DOCUMENT_PATH = os.getenv("DOCUMENT_PATH", "FAQ_DPO_HSE_v3.docx")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
LOG_FILE = os.getenv("LOG_FILE", "questions_log.csv")
MAX_HISTORY = 5  # Количество пар вопрос-ответ в памяти
TELEGRAM_MSG_LIMIT = 4096  # Лимит символов в одном сообщении Telegram

# --- Состояния ---
MENU, WAITING_QUESTION = range(2)

# --- Тексты кнопок ---
BTN_ASK = "❓ Задать вопрос Виртуальному помощнику (24/7)"
BTN_MANAGER = "📞 Связаться с менеджером"
BTN_FAQ = "📋 Часто задаваемые вопросы"
BTN_BACK = "◀️ Назад в меню"

# --- Клавиатуры ---
MAIN_MENU_KEYBOARD = ReplyKeyboardMarkup(
    [
        [KeyboardButton(BTN_ASK)],
        [KeyboardButton(BTN_MANAGER)],
        [KeyboardButton(BTN_FAQ)],
    ],
    resize_keyboard=True,
)

BACK_KEYBOARD = ReplyKeyboardMarkup(
    [[KeyboardButton(BTN_BACK)]],
    resize_keyboard=True,
)

# --- Логирование ---
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)


# --- Загрузка документа ---
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def extract_paragraph_with_links(paragraph, rels) -> str:
    """Извлекает текст параграфа, подставляя URL гиперссылок."""
    result = []
    xml = paragraph._element

    for child in xml:
        tag = etree.QName(child).localname

        if tag == "r":
            texts = child.findall(".//w:t", NSMAP)
            for t in texts:
                if t.text:
                    result.append(t.text)

        elif tag == "hyperlink":
            link_text_parts = []
            for t in child.findall(".//w:t", NSMAP):
                if t.text:
                    link_text_parts.append(t.text)
            link_text = "".join(link_text_parts)

            r_id = child.get(f'{{{NSMAP["r"]}}}id')
            url = ""
            if r_id and r_id in rels:
                url = rels[r_id]

            if url and link_text:
                result.append(f"{link_text} ({url})")
            elif url:
                result.append(url)
            elif link_text:
                result.append(link_text)

    return "".join(result)


def get_rels(doc) -> dict:
    """Извлекает словарь relationship ID -> URL из документа."""
    rels = {}
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            rels[rel.rId] = rel._target
    return rels


def load_document(path: str) -> str:
    """Читает текст из .docx или .txt файла с извлечением гиперссылок."""
    file_path = Path(path)
    if not file_path.exists():
        logger.error(f"Файл не найден: {path}")
        return ""

    if file_path.suffix.lower() == ".docx":
        try:
            doc = Document(str(file_path))
            rels = get_rels(doc)

            paragraphs = []
            for p in doc.paragraphs:
                text = extract_paragraph_with_links(p, rels)
                if text.strip():
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells_text = []
                    for cell in row.cells:
                        cell_parts = []
                        for p in cell.paragraphs:
                            t = extract_paragraph_with_links(p, rels)
                            if t.strip():
                                cell_parts.append(t)
                        if cell_parts:
                            cells_text.append(" ".join(cell_parts))
                    if cells_text:
                        paragraphs.append(" | ".join(cells_text))

            text = "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Ошибка чтения .docx: {e}")
            return ""
    else:
        text = file_path.read_text(encoding="utf-8")

    logger.info(f"Документ загружен: {path} ({len(text)} символов)")
    return text


DOCUMENT_TEXT = load_document(DOCUMENT_PATH)


# --- Логирование вопросов в CSV ---
def init_log_file():
    """Создаёт CSV-файл с заголовками, если он ещё не существует."""
    log_path = Path(LOG_FILE)
    if not log_path.exists():
        with open(log_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([
                "дата_время", "user_id", "username", "имя", "фамилия",
                "вопрос", "ответ", "оценка"
            ])
        logger.info(f"Создан файл лога: {LOG_FILE}")


def log_question(user, question: str, answer: str):
    """Записывает вопрос и ответ в CSV-файл."""
    try:
        with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                user.id,
                user.username or "",
                user.first_name or "",
                user.last_name or "",
                question,
                answer,
                "",  # оценка — заполнится позже
            ])
    except Exception as e:
        logger.error(f"Ошибка записи в лог: {e}")


def update_last_rating(user_id: int, rating: str):
    """Обновляет оценку последнего вопроса пользователя в CSV."""
    try:
        log_path = Path(LOG_FILE)
        if not log_path.exists():
            return

        with open(log_path, "r", newline="", encoding="utf-8") as f:
            rows = list(csv.reader(f))

        # Ищем последнюю строку этого пользователя с пустой оценкой
        for i in range(len(rows) - 1, 0, -1):
            if len(rows[i]) >= 8 and rows[i][1] == str(user_id) and rows[i][7] == "":
                rows[i][7] = rating
                break

        with open(log_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerows(rows)
    except Exception as e:
        logger.error(f"Ошибка обновления оценки: {e}")


init_log_file()

# --- OpenAI клиент ---
client = OpenAI(api_key=OPENAI_API_KEY)

SYSTEM_PROMPT = (
    "Ты — виртуальный помощник программ дополнительного "
    "профессионального образования НИУ ВШЭ. "
    "Отвечай на вопросы ТОЛЬКО на основе предоставленной информации. "
    "Никогда не упоминай слова «документ», «файл», «текст документа» "
    "и подобное. Вместо этого используй фразы вроде «по имеющейся "
    "информации», «по нашим данным», «согласно информации программы». "
    "Если ответа нет в предоставленной информации, скажи: "
    "«К сожалению, у меня нет данных по этому вопросу. "
    "Рекомендую обратиться к менеджеру.» "
    "Отвечай на языке вопроса.\n\n"
    f"--- ИНФОРМАЦИЯ ---\n{DOCUMENT_TEXT}\n--- КОНЕЦ ИНФОРМАЦИИ ---"
)


def ask_question(question: str, history: list) -> str:
    """Отправляет вопрос в OpenAI вместе с контекстом документа и историей диалога."""
    if not DOCUMENT_TEXT:
        return "К сожалению, сейчас я не могу ответить на вопросы. Попробуйте позже или обратитесь к менеджеру."

    try:
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]

        # Добавляем историю диалога
        for h_question, h_answer in history:
            messages.append({"role": "user", "content": h_question})
            messages.append({"role": "assistant", "content": h_answer})

        # Добавляем текущий вопрос
        messages.append({"role": "user", "content": question})

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=messages,
            max_tokens=1024,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Ошибка OpenAI: {e}")
        return f"Произошла ошибка при обработке вопроса: {e}"


# --- Вспомогательные функции ---
def split_message(text: str, limit: int = TELEGRAM_MSG_LIMIT) -> list:
    """Разбивает длинное сообщение на части, не разрывая абзацы."""
    if len(text) <= limit:
        return [text]

    parts = []
    while text:
        if len(text) <= limit:
            parts.append(text)
            break

        # Ищем последний перенос строки в пределах лимита
        cut = text.rfind("\n", 0, limit)
        if cut == -1:
            # Если нет переноса — ищем последний пробел
            cut = text.rfind(" ", 0, limit)
        if cut == -1:
            # Крайний случай — режем по лимиту
            cut = limit

        parts.append(text[:cut])
        text = text[cut:].lstrip("\n")

    return parts


def get_user_name(user) -> str:
    """Возвращает имя пользователя для приветствия."""
    if user.first_name:
        return user.first_name
    if user.username:
        return user.username
    return ""


def get_rating_keyboard() -> InlineKeyboardMarkup:
    """Создаёт inline-клавиатуру для оценки ответа."""
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("👍 Полезно", callback_data="rate_yes"),
            InlineKeyboardButton("👎 Не помогло", callback_data="rate_no"),
        ]
    ])


# --- Обработчики ---
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Команда /start — показываем главное меню."""
    # Очищаем историю при новом старте
    context.user_data["history"] = []

    name = get_user_name(update.effective_user)
    greeting = f"👋 {name}, добро" if name else "👋 Добро"

    await update.message.reply_text(
        f"{greeting} пожаловать в чат-бот программ дополнительного "
        "профессионального образования НИУ ВШЭ!\n\n"
        "Я помогу вам найти ответы на вопросы о поступлении, "
        "программах обучения, стоимости и документах.\n\n"
        "В режиме виртуального помощника я могу рассказать о:\n"
        "• Программах ДПО и форматах обучения\n"
        "• Регистрации и поступлении\n"
        "• Договоре и оферте\n"
        "• Стоимости, оплате и рассрочке\n"
        "• Скидках\n"
        "• Оформлении от юридического лица\n"
        "• Удостоверении и итоговых документах\n"
        "• Платформе обучения (iSpring)\n"
        "• Контактах и налоговом вычете\n\n"
        "Выберите один из пунктов Меню 👇\n\n"
        "📄 Если хотите изучить всю информацию самостоятельно — "
        "ответы на часто задаваемые вопросы вы найдёте внизу веб-страницы: "
        "https://www.hse.ru/edu/dpo/1095672316",
        reply_markup=MAIN_MENU_KEYBOARD,
    )
    return MENU


async def menu_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка нажатий кнопок главного меню."""
    text = update.message.text

    # --- Кнопка 1: Задать вопрос ---
    if text == BTN_ASK:
        # Очищаем историю при входе в режим вопросов
        context.user_data["history"] = []
        await update.message.reply_text(
            "Напишите ваш вопрос, и я постараюсь найти ответ по имеющимся данным.\n\n"
            "Чтобы вернуться в меню, нажмите «◀️ Назад в меню».",
            reply_markup=BACK_KEYBOARD,
        )
        return WAITING_QUESTION

    # --- Кнопка 2: Связаться с менеджером ---
    if text == BTN_MANAGER:
        await update.message.reply_text(
            "Связаться с менеджером можно:\n\n"
            "•  Через Telegram @dposchoolcomm\n"
            "•  По электронной почте: incomm-courses@hse.ru\n"
            "•  По телефону: +7 (495) 772-95-90 (доб. 22390)\n"
            "•  По телефону (альтернативный): +7 (499) 281-65-10\n\n"
            "Для оперативного ответа рекомендуется Telegram "
            "в рабочие часы (будние дни, 10:00–18:00 МСК).",
            reply_markup=MAIN_MENU_KEYBOARD,
        )
        return MENU

    # --- Кнопка 3: Часто задаваемые вопросы ---
    if text == BTN_FAQ:
        await update.message.reply_text(
            "Ответы на часто задаваемые вопросы вы найдёте внизу веб-страницы:\n\n"
            "https://www.hse.ru/edu/dpo/1095672316",
            reply_markup=MAIN_MENU_KEYBOARD,
        )
        return MENU

    # Если текст не совпал ни с одной кнопкой
    await update.message.reply_text(
        "Пожалуйста, выберите один из пунктов меню:",
        reply_markup=MAIN_MENU_KEYBOARD,
    )
    return MENU


async def handle_question(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка вопроса пользователя (состояние WAITING_QUESTION)."""
    text = update.message.text

    # Кнопка «Назад»
    if text == BTN_BACK:
        await update.message.reply_text(
            "Главное меню:",
            reply_markup=MAIN_MENU_KEYBOARD,
        )
        return MENU

    # Получаем историю диалога
    history = context.user_data.get("history", [])

    # Отправляем вопрос в OpenAI
    logger.info(f"Вопрос от {update.effective_user.id}: {text}")
    await update.message.chat.send_action("typing")

    answer = ask_question(text, history)

    # Сохраняем в историю (последние MAX_HISTORY пар)
    history.append((text, answer))
    if len(history) > MAX_HISTORY:
        history = history[-MAX_HISTORY:]
    context.user_data["history"] = history

    # Логируем вопрос и ответ
    log_question(update.effective_user, text, answer)

    # Отправляем ответ (разбиваем если длинный)
    parts = split_message(answer)
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            # Последняя часть — с кнопками оценки
            await update.message.reply_text(part, reply_markup=BACK_KEYBOARD)
            await update.message.reply_text(
                "Был ли ответ полезен?",
                reply_markup=get_rating_keyboard(),
            )
        else:
            await update.message.reply_text(part)

    return WAITING_QUESTION


async def handle_rating(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обработка нажатия кнопки оценки ответа."""
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    rating = "👍 Полезно" if query.data == "rate_yes" else "👎 Не помогло"

    # Обновляем оценку в CSV
    update_last_rating(user_id, rating)

    # Меняем сообщение с кнопками на текст благодарности
    if query.data == "rate_yes":
        await query.edit_message_text("✅ Спасибо за отзыв! Рад, что помог.")
    else:
        await query.edit_message_text(
            "📝 Спасибо за отзыв! Рекомендую обратиться к менеджеру "
            "для получения более подробной информации."
        )


# --- Запуск ---
def main() -> None:
    if not TELEGRAM_TOKEN:
        print("Ошибка: TELEGRAM_TOKEN не задан. Проверьте файл .env")
        return
    if not OPENAI_API_KEY:
        print("Ошибка: OPENAI_API_KEY не задан. Проверьте файл .env")
        return

    app = Application.builder().token(TELEGRAM_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            MENU: [MessageHandler(filters.TEXT & ~filters.COMMAND, menu_handler)],
            WAITING_QUESTION: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_question)],
        },
        fallbacks=[CommandHandler("start", start)],
    )

    app.add_handler(conv_handler)
    # Обработчик оценки — работает независимо от состояния диалога
    app.add_handler(CallbackQueryHandler(handle_rating, pattern="^rate_"))

    logger.info("Бот запущен!")
    app.run_polling()


if __name__ == "__main__":
    main()
