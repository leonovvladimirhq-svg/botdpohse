"""Чат-Бот ДПО Школа коммуникаций НИУ ВШЭ — версия для мессенджера MAX.

Порт Telegram-бота (bot.py) на Bot API мессенджера MAX. Вся продуктовая логика
(база знаний, промпты, вызовы модели, логирование, уведомления админам)
сохранена без изменений; переписан только транспортный слой.

Отличия платформы MAX от Telegram, которые повлияли на код:
  * нет reply-клавиатуры (аналога ReplyKeyboardMarkup) — меню собрано на
    inline-кнопках, которые прикрепляются к сообщениям;
  * нет ConversationHandler — состояние диалога хранится в памяти процесса;
  * лимит сообщения 4000 символов (в Telegram 4096);
  * ответ на нажатие кнопки — POST /answers, он же редактирует сообщение.
"""

import os
import csv
import json
import time
import logging
import threading
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

import requests
from lxml import etree
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

# --- Настройки ---
load_dotenv()

MAX_TOKEN = os.getenv("MAX_TOKEN")
# botapi.max.ru — рабочий домен из Yandex Cloud. У platform-api2.max.ru
# сертификат российского УЦ, которого нет в стандартном ca-certificates,
# поэтому TLS-хендшейк с ВМ падает (unable to get local issuer certificate).
MAX_API_BASE = os.getenv("MAX_API_BASE", "https://botapi.max.ru").rstrip("/")
YANDEX_API_KEY = os.getenv("YANDEX_API_KEY")
YANDEX_BASE_URL = os.getenv("YANDEX_BASE_URL", "https://llm.api.cloud.yandex.net/v1")
DOCUMENT_PATH = os.getenv("DOCUMENT_PATH", "FAQ_DPO_HSE_v5.docx")
MODEL_URI = os.getenv("MODEL_URI")
LOG_FILE = os.getenv("LOG_FILE", "questions_log.csv")
ADMIN_CHAT_ID = os.getenv("ADMIN_CHAT_ID", "")  # MAX user_id администратора
ADMIN_CHAT_ID_2 = os.getenv("ADMIN_CHAT_ID_2", "")  # MAX user_id второго администратора
MAX_HISTORY = 5  # Количество пар вопрос-ответ в памяти
MAX_MSG_LIMIT = 3900  # Лимит символов в одном сообщении MAX (жёсткий лимит API — 4000)
POLL_TIMEOUT = 30  # Long polling: сколько секунд сервер держит соединение
WORKERS = 4  # Одновременно обрабатываемых обновлений

# --- Состояния ---
MENU, WAITING_QUESTION = range(2)

# --- Тексты кнопок ---
BTN_ASK = "❓ Задать вопрос Виртуальному помощнику (24/7)"
BTN_MANAGER = "📞 Связаться с менеджером"
BTN_FAQ = "📋 Часто задаваемые вопросы"
BTN_BACK = "◀️ Назад в меню"

# --- payload'ы callback-кнопок ---
CB_ASK = "menu_ask"
CB_MANAGER = "menu_manager"
CB_FAQ = "menu_faq"
CB_BACK = "menu_back"
CB_RATE_YES = "rate_yes"
CB_RATE_NO = "rate_no"

# --- Контакты менеджеров для сценария, когда бот не смог ответить ---
MANAGER_PHONES_TEXT = (
    "📞 Телефон менеджеров программ:\n"
    "•  +7 (495) 772-95-90 (доб. 23644)"
)

# --- Логирование ---
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)


# --- Клавиатуры (в MAX — только inline) ---
def kb(rows: list) -> list:
    """Оборачивает строки кнопок во вложение inline_keyboard."""
    return [{"type": "inline_keyboard", "payload": {"buttons": rows}}]


def callback_button(text: str, payload: str) -> dict:
    return {"type": "callback", "text": text, "payload": payload}


MAIN_MENU_KEYBOARD = kb([
    [callback_button(BTN_ASK, CB_ASK)],
    [callback_button(BTN_MANAGER, CB_MANAGER)],
    [callback_button(BTN_FAQ, CB_FAQ)],
])

BACK_KEYBOARD = kb([[callback_button(BTN_BACK, CB_BACK)]])

RATING_KEYBOARD = kb([[
    callback_button("👍 Полезно", CB_RATE_YES),
    callback_button("👎 Не помогло", CB_RATE_NO),
]])


# --- Загрузка документа ---
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def extract_paragraph_with_links(paragraph, rels) -> str:
    """Извлекает текст параграфа, подставляя URL гиперссылок."""
    result = []
    xml = paragraph._element

    for child in xml:
        tag = etree.QName(child).localname

        if tag == "r":
            texts = child.findall(".//w:t", NSMAP)
            for t in texts:
                if t.text:
                    result.append(t.text)

        elif tag == "hyperlink":
            link_text_parts = []
            for t in child.findall(".//w:t", NSMAP):
                if t.text:
                    link_text_parts.append(t.text)
            link_text = "".join(link_text_parts)

            r_id = child.get(f'{{{NSMAP["r"]}}}id')
            url = ""
            if r_id and r_id in rels:
                url = rels[r_id]

            if url and link_text:
                result.append(f"{link_text} ({url})")
            elif url:
                result.append(url)
            elif link_text:
                result.append(link_text)

    return "".join(result)


def get_rels(doc) -> dict:
    """Извлекает словарь relationship ID -> URL из документа."""
    rels = {}
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            rels[rel.rId] = rel._target
    return rels


def load_document(path: str) -> str:
    """Читает текст из .docx или .txt файла с извлечением гиперссылок."""
    file_path = Path(path)
    if not file_path.exists():
        logger.error(f"Файл не найден: {path}")
        return ""

    if file_path.suffix.lower() == ".docx":
        try:
            doc = Document(str(file_path))
            rels = get_rels(doc)

            paragraphs = []
            for p in doc.paragraphs:
                text = extract_paragraph_with_links(p, rels)
                if text.strip():
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells_text = []
                    for cell in row.cells:
                        cell_parts = []
                        for p in cell.paragraphs:
                            t = extract_paragraph_with_links(p, rels)
                            if t.strip():
                                cell_parts.append(t)
                        if cell_parts:
                            cells_text.append(" ".join(cell_parts))
                    if cells_text:
                        paragraphs.append(" | ".join(cells_text))

            text = "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Ошибка чтения .docx: {e}")
            return ""
    else:
        text = file_path.read_text(encoding="utf-8")

    logger.info(f"Документ загружен: {path} ({len(text)} символов)")
    return text


DOCUMENT_TEXT = load_document(DOCUMENT_PATH)


# --- Логирование вопросов в CSV ---
_log_lock = threading.Lock()


def init_log_file():
    """Создаёт CSV-файл с заголовками, если он ещё не существует."""
    log_path = Path(LOG_FILE)
    if not log_path.exists():
        with open(log_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([
                "дата_время", "user_id", "username", "имя", "фамилия",
                "вопрос", "ответ", "оценка"
            ])
        logger.info(f"Создан файл лога: {LOG_FILE}")


def log_question(user: dict, question: str, answer: str):
    """Записывает вопрос и ответ в CSV-файл."""
    try:
        with _log_lock, open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                user.get("user_id", ""),
                user.get("username") or "",
                user.get("first_name") or "",
                user.get("last_name") or "",
                question,
                answer,
                "",  # оценка — заполнится позже
            ])
    except Exception as e:
        logger.error(f"Ошибка записи в лог: {e}")


def update_last_rating(user_id: int, rating: str):
    """Обновляет оценку последнего вопроса пользователя в CSV."""
    try:
        log_path = Path(LOG_FILE)
        if not log_path.exists():
            return

        with _log_lock:
            with open(log_path, "r", newline="", encoding="utf-8") as f:
                rows = list(csv.reader(f))

            # Ищем последнюю строку этого пользователя с пустой оценкой
            for i in range(len(rows) - 1, 0, -1):
                if len(rows[i]) >= 8 and rows[i][1] == str(user_id) and rows[i][7] == "":
                    rows[i][7] = rating
                    break

            with open(log_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerows(rows)
    except Exception as e:
        logger.error(f"Ошибка обновления оценки: {e}")


init_log_file()


# --- Клиент MAX Bot API ---
class MaxBot:
    """Минимальный синхронный клиент Bot API мессенджера MAX.

    Реализованы только методы, нужные боту: long polling, отправка и
    редактирование сообщений, ответ на callback, индикатор «печатает».
    """

    def __init__(self, token: str, base_url: str = MAX_API_BASE):
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        # Токен передаётся заголовком: query-параметр access_token объявлен
        # устаревшим и отдаёт 401.
        self.session.headers.update({"Authorization": token})

    def _request(self, method: str, path: str, *, params=None, json_body=None,
                 timeout: int = 30):
        url = f"{self.base_url}{path}"
        resp = self.session.request(
            method, url, params=params, json=json_body, timeout=timeout
        )
        if resp.status_code != 200:
            raise RuntimeError(
                f"MAX API {method} {path} -> {resp.status_code}: {resp.text[:300]}"
            )
        return resp.json()

    def get_me(self) -> dict:
        return self._request("GET", "/me")

    def get_updates(self, marker=None, limit: int = 100) -> dict:
        params = {"timeout": POLL_TIMEOUT, "limit": limit}
        if marker is not None:
            params["marker"] = marker
        # Читаем дольше, чем длится long polling, иначе рвём соединение сами.
        return self._request(
            "GET", "/updates", params=params, timeout=POLL_TIMEOUT + 30
        )

    def send_message(self, user_id: int, text: str, attachments=None,
                     notify: bool = True) -> dict:
        body = {"text": text, "notify": notify}
        if attachments:
            body["attachments"] = attachments
        return self._request(
            "POST", "/messages", params={"user_id": user_id}, json_body=body
        )

    def answer_callback(self, callback_id: str, notification: str = None,
                        message: dict = None) -> dict:
        """Отвечает на нажатие кнопки.

        message заменяет содержимое сообщения, к которому была прикреплена
        клавиатура (аналог edit_message_text в Telegram).
        """
        body = {}
        if notification:
            body["notification"] = notification
        if message is not None:
            body["message"] = message
        return self._request(
            "POST", "/answers", params={"callback_id": callback_id}, json_body=body
        )

    def send_action(self, chat_id: int, action: str = "typing_on") -> None:
        """Индикатор «печатает». Ошибки не критичны — просто логируем."""
        try:
            self._request(
                "POST", f"/chats/{chat_id}/actions", json_body={"action": action}
            )
        except Exception as e:
            logger.debug(f"send_action не удался: {e}")

    def set_commands(self, commands: list) -> None:
        """Регистрирует список команд бота (меню команд в клиенте MAX)."""
        try:
            self._request("PATCH", "/me", json_body={"commands": commands})
        except Exception as e:
            logger.warning(f"Не удалось задать команды бота: {e}")


bot = MaxBot(MAX_TOKEN or "")


# --- Уведомления администраторам ---
def notify_admins(text: str):
    """Дублирует событие администраторам в MAX.

    ADMIN_CHAT_ID — это MAX user_id (не Telegram!). Узнать свой ID можно,
    отправив боту команду /whoami. Бот может писать только тем, кто уже
    начал с ним диалог.
    """
    for admin_id in (ADMIN_CHAT_ID, ADMIN_CHAT_ID_2):
        if not admin_id:
            continue
        try:
            for part in split_message(text):
                bot.send_message(int(admin_id), part)
            logger.info(f"Уведомление отправлено админу {admin_id}")
        except Exception as e:
            logger.error(f"Ошибка отправки уведомления админу {admin_id}: {e}")


# --- Клиент Yandex AI Studio (OpenAI-совместимый эндпоинт) ---
client = OpenAI(api_key=YANDEX_API_KEY, base_url=YANDEX_BASE_URL)

SYSTEM_PROMPT = (
    "Ты — виртуальный помощник программ дополнительного "
    "профессионального образования НИУ ВШЭ. "
    "Отвечай на вопросы ТОЛЬКО на основе предоставленной информации. "
    "Никогда не упоминай слова «документ», «файл», «текст документа» "
    "и подобное. Вместо этого используй фразы вроде «по имеющейся "
    "информации», «по нашим данным», «согласно информации программы». "
    "Если ответа нет в предоставленной информации, скажи: "
    "«К сожалению, у меня нет данных по этому вопросу. "
    "Рекомендую обратиться к менеджеру.» "
    "Отвечай на языке вопроса. "
    "ВАЖНО про оформление: пиши обычным текстом без любой markdown-разметки. "
    "Не используй звёздочки (*, **), подчёркивания (_, __), решётки (#), "
    "обратные кавычки (`), угловые скобки для тегов. Если нужно что-то "
    "выделить — используй обычные слова или КАПС. Для списков допустимы "
    "только дефисы (- ) или цифры (1. ). Не вставляй пустые строки между "
    "пунктами без необходимости.\n\n"
    f"--- ИНФОРМАЦИЯ ---\n{DOCUMENT_TEXT}\n--- КОНЕЦ ИНФОРМАЦИИ ---"
)


def ask_question(question: str, history: list) -> str:
    """Отправляет вопрос в модель вместе с контекстом документа и историей диалога."""
    if not DOCUMENT_TEXT:
        return "К сожалению, сейчас я не могу ответить на вопросы. Попробуйте позже или обратитесь к менеджеру."

    try:
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]

        # Добавляем историю диалога
        for h_question, h_answer in history:
            messages.append({"role": "user", "content": h_question})
            messages.append({"role": "assistant", "content": h_answer})

        # Добавляем текущий вопрос
        messages.append({"role": "user", "content": question})

        response = client.chat.completions.create(
            model=MODEL_URI,
            messages=messages,
            max_completion_tokens=1024,
            extra_body={"reasoning_effort": "none"},
        )
        result = response.choices[0].message.content
        return strip_markdown(result) if result else ""
    except Exception as e:
        logger.error(f"Ошибка обращения к модели: {e}")
        return f"Произошла ошибка при обработке вопроса: {e}"


def strip_markdown(text: str) -> str:
    """Удаляет markdown-разметку, которую модель иногда добавляет в ответ.

    Бот шлёт сообщения как plain text (без format), поэтому **жирный**
    и _курсив_ отображаются у пользователя как мусор со звёздочками.
    """
    if not text:
        return text
    import re
    # **bold** / __bold__  -> bold
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"__(.+?)__", r"\1", text, flags=re.DOTALL)
    # *italic* / _italic_  -> italic (только парные, не задевая отдельные * в тексте)
    text = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_(?!_)([^_\n]+?)(?<!_)_(?!_)", r"\1", text)
    # `code` -> code
    text = re.sub(r"`([^`\n]+?)`", r"\1", text)
    # ### headings -> plain
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    return text


# --- Детекция «нет ответа» и подбор переформулировок ---
NO_DATA_MARKER = "нет данных по этому вопросу"  # фраза из SYSTEM_PROMPT


def is_no_data_answer(answer: str) -> bool:
    """Возвращает True, если ответ пустой или содержит маркер «нет данных»."""
    if not answer or not answer.strip():
        return True
    return NO_DATA_MARKER.lower() in answer.lower()


def extract_faq_questions(doc_text: str) -> list:
    """Извлекает список вопросов из FAQ.

    В исходном документе каждый вопрос начинается с префикса "В: " и
    заканчивается на "?". Возвращаем чистые формулировки без префикса.
    """
    questions = []
    for line in doc_text.split("\n"):
        s = line.strip()
        if s.startswith("В:"):
            q = s[2:].lstrip(": ").strip()
            if q:
                questions.append(q)
    return questions


FAQ_QUESTIONS = extract_faq_questions(DOCUMENT_TEXT)
logger.info(f"Извлечено вопросов из FAQ: {len(FAQ_QUESTIONS)}")


def _build_suggest_messages(question: str) -> list:
    """Готовит messages для модели: список FAQ-вопросов + вопрос пользователя."""
    numbered = "\n".join(f"{i + 1}. {q}" for i, q in enumerate(FAQ_QUESTIONS))
    system_prompt = (
        "Ты помогаешь пользователю переформулировать вопрос о программах "
        "дополнительного профессионального образования НИУ ВШЭ (Школа коммуникаций). "
        "Ниже — список вопросов из нашего FAQ, на каждый из которых у нас есть ответ. "
        "Пользователь задал свой вопрос — возможно, нечётко или неточно. "
        "Твоя задача — подобрать из списка от 1 до 3 вопросов, которые ближе всего "
        "по смыслу к запросу пользователя. Будь готов предлагать варианты, даже "
        "если совпадение по теме приблизительное — пользователю важнее увидеть "
        "близкие темы, чем услышать «не нашлось». "
        "Возвращай пустой список ТОЛЬКО если запрос пользователя совершенно не "
        "относится к темам обучения, поступления, оплаты, документов, платформы "
        "и других тем из FAQ (например, спрашивают о погоде или политике). "
        "Не придумывай новые вопросы — выбирай только из списка ниже, точно как они "
        "там записаны. Ответ строго в JSON виде "
        "{\"suggestions\": [\"...\", \"...\", \"...\"]}. Без какого-либо текста вне JSON.\n\n"
        "--- СПИСОК ВОПРОСОВ FAQ ---\n"
        f"{numbered}\n"
        "--- КОНЕЦ СПИСКА ---"
    )
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": question},
    ]


def suggest_reformulations(question: str, max_items: int = 3) -> list:
    """Подбирает до max_items похожих вопросов из FAQ через модель."""
    if not FAQ_QUESTIONS:
        return []
    try:
        response = client.chat.completions.create(
            model=MODEL_URI,
            messages=_build_suggest_messages(question),
            max_completion_tokens=1024,
            response_format={"type": "json_object"},
            extra_body={"reasoning_effort": "none"},
        )
        raw = response.choices[0].message.content or "{}"
        finish = response.choices[0].finish_reason
        data = json.loads(raw)
        items = data.get("suggestions", []) or []
        cleaned = [s.strip() for s in items if isinstance(s, str) and s.strip()]
        # Оставляем только те, которые реально есть в FAQ
        known = {q.lower(): q for q in FAQ_QUESTIONS}
        filtered = []
        for s in cleaned:
            if s.lower() in known:
                filtered.append(known[s.lower()])
            else:
                # Если модель чуть переформулировала — попробуем найти ближайший по
                # подстрочному совпадению (упрощённо)
                for q in FAQ_QUESTIONS:
                    if s.lower() in q.lower() or q.lower() in s.lower():
                        filtered.append(q)
                        break
        # Уникальные, с сохранением порядка
        seen = set()
        unique = []
        for q in filtered:
            if q not in seen:
                seen.add(q)
                unique.append(q)
        logger.info(
            f"suggest_reformulations: finish={finish}, "
            f"raw_count={len(cleaned)}, kept={len(unique)}"
        )
        return unique[:max_items]
    except Exception as e:
        logger.error(f"Ошибка suggest_reformulations: {e}")
        return []


# --- Вспомогательные функции ---
def split_message(text: str, limit: int = MAX_MSG_LIMIT) -> list:
    """Разбивает длинное сообщение на части, не разрывая абзацы."""
    if len(text) <= limit:
        return [text]

    parts = []
    while text:
        if len(text) <= limit:
            parts.append(text)
            break

        # Ищем последний перенос строки в пределах лимита
        cut = text.rfind("\n", 0, limit)
        if cut == -1:
            # Если нет переноса — ищем последний пробел
            cut = text.rfind(" ", 0, limit)
        if cut == -1:
            # Крайний случай — режем по лимиту
            cut = limit

        parts.append(text[:cut])
        text = text[cut:].lstrip("\n")

    return parts


def get_user_name(user: dict) -> str:
    """Возвращает имя пользователя для приветствия."""
    if user.get("first_name"):
        return user["first_name"]
    if user.get("username"):
        return user["username"]
    return ""


def describe_user(user: dict) -> str:
    """Строка «Имя Фамилия (@username)» для уведомлений админам."""
    username = f"@{user['username']}" if user.get("username") else "нет"
    name = f"{user.get('first_name') or ''} {user.get('last_name') or ''}".strip() or "нет"
    return f"{name} ({username})"


# --- Состояние диалогов (аналог ConversationHandler + user_data) ---
_states = {}  # user_id -> {"state": int, "history": list, "chat_id": int}
_states_lock = threading.Lock()


def get_state(user_id: int) -> dict:
    with _states_lock:
        return _states.setdefault(
            user_id, {"state": MENU, "history": [], "chat_id": None}
        )


def reset_state(user_id: int):
    with _states_lock:
        _states.pop(user_id, None)


# --- Обработчики ---
_last_welcome = {}  # user_id -> время последнего приветствия (для дедупликации)


def send_welcome(user_id: int, user: dict):
    """Команда /start — показываем главное меню."""
    st = get_state(user_id)
    st["state"] = MENU
    st["history"] = []
    _last_welcome[user_id] = time.monotonic()

    name = get_user_name(user)
    greeting = f"👋 {name}, добро" if name else "👋 Добро"

    bot.send_message(
        user_id,
        f"{greeting} пожаловать в чат-бот программ дополнительного "
        "профессионального образования НИУ ВШЭ!\n\n"
        "Я помогу вам найти ответы на вопросы о поступлении, "
        "программах обучения, стоимости и документах.\n\n"
        "В режиме виртуального помощника я могу рассказать о:\n"
        "• Программах ДПО и форматах обучения\n"
        "• Регистрации и поступлении\n"
        "• Договоре и оферте\n"
        "• Стоимости, оплате и рассрочке\n"
        "• Скидках\n"
        "• Оформлении от юридического лица\n"
        "• Удостоверении и итоговых документах\n"
        "• Платформе обучения (iSpring)\n"
        "• Контактах и налоговом вычете\n\n"
        "Выберите один из пунктов Меню 👇\n\n"
        "📄 Если хотите изучить всю информацию самостоятельно — "
        "полный документ с ответами на часто задаваемые вопросы доступен по ссылке: "
        "https://disk.360.yandex.ru/i/tNTbuVq6Bp385A",
        attachments=MAIN_MENU_KEYBOARD,
    )


def show_menu(user_id: int):
    """Возврат в меню без длинного приветствия."""
    st = get_state(user_id)
    st["state"] = MENU
    bot.send_message(
        user_id,
        "Выберите один из пунктов меню:",
        attachments=MAIN_MENU_KEYBOARD,
    )


def handle_menu_choice(user_id: int, payload: str):
    """Обработка нажатий кнопок главного меню."""
    st = get_state(user_id)

    # --- Кнопка 1: Задать вопрос ---
    if payload == CB_ASK:
        # Очищаем историю при входе в режим вопросов
        st["history"] = []
        st["state"] = WAITING_QUESTION
        bot.send_message(
            user_id,
            "Напишите ваш вопрос, и я постараюсь найти ответ по имеющимся данным.\n\n"
            "Чтобы вернуться в меню, нажмите «◀️ Назад в меню».",
            attachments=BACK_KEYBOARD,
        )
        return

    # --- Кнопка 2: Связаться с менеджером ---
    if payload == CB_MANAGER:
        st["state"] = MENU
        bot.send_message(
            user_id,
            "Связаться с менеджером можно:\n\n"
            "•  Через Telegram @dposchoolcomm\n"
            "•  По электронной почте: incomm-courses@hse.ru\n"
            "•  По телефону: +7 (495) 772-95-90 (доб. 23644)\n\n"
            "Для оперативного ответа рекомендуется Telegram "
            "в рабочие часы (будние дни, 10:00–18:00 МСК).",
            attachments=MAIN_MENU_KEYBOARD,
        )
        return

    # --- Кнопка 3: Часто задаваемые вопросы ---
    if payload == CB_FAQ:
        st["state"] = MENU
        bot.send_message(
            user_id,
            "Полный документ с ответами на часто задаваемые вопросы доступен по ссылке:\n\n"
            "https://disk.360.yandex.ru/i/tNTbuVq6Bp385A",
            attachments=MAIN_MENU_KEYBOARD,
        )
        return


def handle_question(user_id: int, chat_id: int, user: dict, text: str):
    """Обработка вопроса пользователя (состояние WAITING_QUESTION)."""
    st = get_state(user_id)
    history = st.get("history", [])

    logger.info(f"Вопрос от {user_id}: {text}")
    if chat_id:
        bot.send_action(chat_id, "typing_on")

    answer = ask_question(text, history)

    # Если ответа нет — пробуем подобрать похожие вопросы из FAQ
    if is_no_data_answer(answer):
        suggestions = suggest_reformulations(text, max_items=3)
        if suggestions:
            bullets = "\n".join(f"•  {s}" for s in suggestions)
            answer = (
                "Возможно, вы хотели спросить:\n"
                f"{bullets}\n\n"
                f"{MANAGER_PHONES_TEXT}"
            )
        else:
            answer = (
                "К сожалению, не удалось получить ответ. "
                "Попробуйте переформулировать вопрос или обратитесь к менеджеру.\n\n"
                f"{MANAGER_PHONES_TEXT}"
            )

    # Сохраняем в историю (последние MAX_HISTORY пар)
    history.append((text, answer))
    if len(history) > MAX_HISTORY:
        history = history[-MAX_HISTORY:]
    st["history"] = history

    # Логируем вопрос и ответ
    log_question(user, text, answer)

    # Уведомляем администраторов
    notify_admins(
        f"📩 Новый вопрос\n\n"
        f"👤 Пользователь: {describe_user(user)}\n"
        f"🆔 ID: {user_id}\n\n"
        f"❓ Вопрос:\n{text}\n\n"
        f"💬 Ответ:\n{answer[:3000]}"
    )

    # Отправляем ответ (разбиваем если длинный)
    parts = split_message(answer)
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            # Последняя часть — с кнопкой возврата в меню
            bot.send_message(user_id, part, attachments=BACK_KEYBOARD)
            bot.send_message(
                user_id,
                "Был ли ответ полезен?",
                attachments=RATING_KEYBOARD,
            )
        else:
            bot.send_message(user_id, part)


def handle_rating(user_id: int, user: dict, payload: str, callback_id: str):
    """Обработка нажатия кнопки оценки ответа."""
    rating = "👍 Полезно" if payload == CB_RATE_YES else "👎 Не помогло"

    # Обновляем оценку в CSV
    update_last_rating(user_id, rating)

    # Меняем сообщение с кнопками на текст благодарности. Делаем это первым
    # делом: пока не ответили на callback, у пользователя крутится индикатор
    # на кнопке, а рассылка админам может быть небыстрой.
    if payload == CB_RATE_YES:
        new_text = "✅ Спасибо за отзыв! Рад, что помог."
    else:
        new_text = (
            "📝 Спасибо за отзыв! Рекомендую обратиться к менеджеру "
            "для получения более подробной информации."
        )
    bot.answer_callback(
        callback_id,
        message={"text": new_text, "attachments": []},
    )

    # Уведомляем администраторов об оценке
    username = f"@{user['username']}" if user.get("username") else "нет"
    notify_admins(f"⭐ Оценка: {rating}\n👤 От: {username} (ID: {user_id})")


# --- Диспетчеризация обновлений ---
def process_update(update: dict):
    """Разбирает одно обновление и вызывает нужный обработчик."""
    utype = update.get("update_type")

    # Пользователь открыл диалог с ботом (аналог первого /start в Telegram)
    if utype == "bot_started":
        user = update.get("user") or {}
        user_id = user.get("user_id")
        if not user_id:
            return
        get_state(user_id)["chat_id"] = update.get("chat_id")
        logger.info(f"bot_started от {user_id} ({describe_user(user)})")
        send_welcome(user_id, user)
        return

    # Нажата inline-кнопка
    if utype == "message_callback":
        callback = update.get("callback") or {}
        user = callback.get("user") or {}
        user_id = user.get("user_id")
        payload = callback.get("payload")
        callback_id = callback.get("callback_id")
        if not user_id or not callback_id:
            return

        message = update.get("message") or {}
        chat_id = (message.get("recipient") or {}).get("chat_id")
        if chat_id:
            get_state(user_id)["chat_id"] = chat_id

        if payload in (CB_RATE_YES, CB_RATE_NO):
            handle_rating(user_id, user, payload, callback_id)
            return

        # Остальные кнопки: сначала закрываем «часики» на кнопке
        try:
            bot.answer_callback(callback_id)
        except Exception as e:
            logger.debug(f"answer_callback не удался: {e}")

        if payload == CB_BACK:
            # Кнопка «Назад» — полный рестарт диалога, как в Telegram-версии
            reset_state(user_id)
            get_state(user_id)["chat_id"] = chat_id
            send_welcome(user_id, user)
            return

        if payload in (CB_ASK, CB_MANAGER, CB_FAQ):
            handle_menu_choice(user_id, payload)
            return

        logger.warning(f"Неизвестный payload кнопки: {payload}")
        return

    # Обычное текстовое сообщение
    if utype == "message_created":
        message = update.get("message") or {}
        user = message.get("sender") or {}
        user_id = user.get("user_id")
        if not user_id or user.get("is_bot"):
            return

        recipient = message.get("recipient") or {}
        # Работаем только в личных диалогах, как и Telegram-версия
        if recipient.get("chat_type") not in (None, "dialog"):
            return
        chat_id = recipient.get("chat_id")

        body = message.get("body") or {}
        text = (body.get("text") or "").strip()
        if not text:
            return

        st = get_state(user_id)
        st["chat_id"] = chat_id

        # Команды
        if text.startswith("/"):
            command = text.split()[0].split("@")[0].lower()
            if command == "/start":
                # При открытии диалога MAX может прислать и bot_started,
                # и сообщение «/start» — второе приветствие подряд гасим.
                last = _last_welcome.get(user_id)
                if last is not None and time.monotonic() - last < 5:
                    logger.info(f"Повторный /start от {user_id} пропущен")
                    return
                send_welcome(user_id, user)
                return
            if command == "/whoami":
                # Служебная команда: помогает администраторам узнать свой MAX ID
                # для ADMIN_CHAT_ID (в MAX другие ID, чем в Telegram).
                bot.send_message(
                    user_id,
                    f"Ваш MAX ID: {user_id}\nchat_id: {chat_id}",
                )
                return
            # Неизвестная команда — ведём себя как при обычном тексте
            show_menu(user_id)
            return

        # Кнопки MAX присылают callback, но пользователь мог набрать текст руками
        if text == BTN_BACK:
            reset_state(user_id)
            get_state(user_id)["chat_id"] = chat_id
            send_welcome(user_id, user)
            return

        if st["state"] == WAITING_QUESTION:
            handle_question(user_id, chat_id, user, text)
        else:
            # В меню бот ждёт нажатия кнопки
            if text == BTN_ASK:
                handle_menu_choice(user_id, CB_ASK)
            elif text == BTN_MANAGER:
                handle_menu_choice(user_id, CB_MANAGER)
            elif text == BTN_FAQ:
                handle_menu_choice(user_id, CB_FAQ)
            else:
                show_menu(user_id)
        return

    # Пользователь удалил диалог или остановил бота — чистим состояние
    if utype in ("bot_stopped", "dialog_removed", "dialog_cleared"):
        user = update.get("user") or {}
        if user.get("user_id"):
            reset_state(user["user_id"])
        return

    logger.debug(f"Необработанный тип обновления: {utype}")


def safe_process(update: dict):
    try:
        process_update(update)
    except Exception as e:
        logger.exception(f"Ошибка обработки обновления: {e}")


# --- Запуск ---
def main() -> None:
    if not MAX_TOKEN:
        print("Ошибка: MAX_TOKEN не задан. Проверьте файл .env")
        return
    if not YANDEX_API_KEY:
        print("Ошибка: YANDEX_API_KEY не задан. Проверьте файл .env")
        return
    if not MODEL_URI:
        print("Ошибка: MODEL_URI не задан. Проверьте файл .env")
        return

    me = bot.get_me()
    logger.info(
        f"Бот запущен: {me.get('name')} (@{me.get('username')}, id={me.get('user_id')})"
    )

    bot.set_commands([
        {"name": "start", "description": "Главное меню"},
    ])

    pool = ThreadPoolExecutor(max_workers=WORKERS)
    marker = None
    backoff = 1

    while True:
        try:
            data = bot.get_updates(marker=marker)
            backoff = 1
        except Exception as e:
            logger.error(f"Ошибка получения обновлений: {e}; повтор через {backoff} с")
            time.sleep(backoff)
            backoff = min(backoff * 2, 60)
            continue

        updates = data.get("updates") or []
        # marker двигаем всегда, даже если событий не было
        marker = data.get("marker", marker)

        for update in updates:
            pool.submit(safe_process, update)


if __name__ == "__main__":
    main()
